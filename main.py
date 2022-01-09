import random
import datetime
from datetime import date
from tkinter import *
from tkinter import ttk, filedialog
import os
import subprocess
import re

from docx import Document
from docxtpl import DocxTemplate

root = Tk()
filename = os.getcwd() + '/КАРТОЧКА.docx'
lst = []
new_filename = os.getcwd() + '/ПОЛУЧИЛОСЬ.docx'


def get_table():
    table.delete(*table.get_children())
    num = numField.get()
    for i in range(int(num)):
        start_date = date.today().replace(day=1, month=1).toordinal()
        end_date = date.today().toordinal()
        random_day = date.fromordinal(random.randint(start_date, end_date))
        lst.append((f'товар{i}', random.randint(100, 100000),
                    random_day.strftime('%d.%m.%Y'),
                    random.randint(100, 1000) / 100,
                    random.randint(1, 100)))
    for row in lst:
        table.insert('', END, values=row)


def func1():
    card = field1.get()  # int
    org = field2.get()  # str
    sub_org = field3.get()  # str
    card_date = field4.get()  # date
    card_name = field5.get()  # str
    # subprocess.call(['open', filename])
    doc = DocxTemplate(filename)
    text = ''
    for row in lst:
        for item in row:
            text += str(item) + '\t'

    lst.clear()
    for line in table.get_children():
        lst.append(table.item(line)['values'])

    context = {'card': int(card),
               'org': str(org),
               'sub_org': str(sub_org),
               'card_date': datetime.datetime.strptime(card_date, '%d.%m.%Y').strftime('%d.%m.%Y'),
               'card_name': str(card_name),
               'tbl_contents': lst
               }
    doc.render(context)
    doc.save(new_filename)
    subprocess.call(['open', new_filename])


def func2():
    global filename
    root.withdraw()
    filename = filedialog.askopenfilename(initialdir=os.getcwd(), filetypes=[('Word files', '.docx .dotx')])
    print(filename)
    root.deiconify()


def func3():
    doc_obj = Document(new_filename)
    regex = re.compile(r"\.")
    replace = r','
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for t in doc_obj.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if regex.search(p.text):
                        inline = p.runs
                        # Loop added to work with runs (strings with same style)
                        for i in range(len(inline)):
                            if regex.search(inline[i].text):
                                text = regex.sub(replace, inline[i].text)
                                inline[i].text = text

    doc_obj.save(new_filename)
    subprocess.call(['open', new_filename])


def treeview_sort_column(treeview: ttk.Treeview, col, reverse: bool):
    """
    to sort the table by column when clicking in column
    """
    try:
        data_list = [
            (int(treeview.set(k, col)), k) for k in treeview.get_children("")
        ]
    except Exception:
        data_list = [(treeview.set(k, col), k) for k in treeview.get_children("")]

    try:
        data_list.sort(key=lambda x: datetime.datetime.strptime(x[0], '%d.%m.%Y'), reverse=reverse)
    except Exception:
        data_list.sort(reverse=reverse)

    # rearrange items in sorted positions
    for index, (val, k) in enumerate(data_list):
        treeview.move(k, "", index)

    # reverse sort next time
    treeview.heading(
        column=col,
        text=col,
        command=lambda _col=col: treeview_sort_column(
            treeview, _col, not reverse
        ),
    )


root['bg'] = 'black'
root.title('Павловец Ангелина Максимовна, 2021, 4 курс, 4 группа')
root.geometry('1300x700')

# номер карточки
frame_first = Frame(root, bg='black', bd=3)
frame_first.place(relx=0.05, rely=0.05, relwidth=0.2, relheight=0.05)
frame_first_text = Frame(root, bg='black', bd=3)
frame_first_text.place(relx=0.05, rely=0, relwidth=0.2, relheight=0.05)
info1 = Label(frame_first_text, text='Введите номер карточки:', bg='black', font=40)
info1.pack()
field1 = Entry(frame_first, bg='gray', font=30)
field1.pack()

# организация
frame_second = Frame(root, bg='black', bd=3)
frame_second.place(relx=0.05, rely=0.15, relwidth=0.2, relheight=0.05)
frame_second_text = Frame(root, bg='black', bd=3)
frame_second_text.place(relx=0.05, rely=0.1, relwidth=0.2, relheight=0.05)
info2 = Label(frame_second_text, text='Организация:', bg='black', font=40)
info2.pack()
field2 = Entry(frame_second, bg='gray', font=30)
field2.pack()

# подразделение
frame_third = Frame(root, bg='black', bd=3)
frame_third.place(relx=0.05, rely=0.25, relwidth=0.2, relheight=0.05)
frame_third_text = Frame(root, bg='black', bd=3)
frame_third_text.place(relx=0.05, rely=0.2, relwidth=0.2, relheight=0.05)
info3 = Label(frame_third_text, text='Структурное подразделение:', bg='black', font=40)
info3.pack()
field3 = Entry(frame_third, bg='gray', font=30)
field3.pack()

# дата
frame_forth = Frame(root, bg='black', bd=3)
frame_forth.place(relx=0.05, rely=0.35, relwidth=0.2, relheight=0.05)
frame_forth_text = Frame(root, bg='black', bd=3)
frame_forth_text.place(relx=0.05, rely=0.3, relwidth=0.2, relheight=0.05)
info4 = Label(frame_forth_text, text='Дата:', bg='black', font=40)
info4.pack()
field4 = Entry(frame_forth, bg='gray', font=30)
field4.pack()

# подпись
frame_fifth = Frame(root, bg='black', bd=3)
frame_fifth.place(relx=0.05, rely=0.45, relwidth=0.2, relheight=0.05)
frame_fifth_text = Frame(root, bg='black', bd=3)
frame_fifth_text.place(relx=0.05, rely=0.4, relwidth=0.2, relheight=0.05)
info5 = Label(frame_fifth_text, text='Подпись:', bg='black', font=40)
info5.pack()
field5 = Entry(frame_fifth, bg='gray', font=30)
field5.pack()

# табличка
frame_list = Frame(root, bg='white', bd=3)
frame_list.place(relx=0.35, rely=0.05, relwidth=0.6, relheight=0.9)
heads = ['Наименование товара', 'Артикул', 'Дата последнего поступления', 'Стоимость единицы товара', 'Количество']
anch = ['w', 'e', 'center', 'e', 'e']
s_ = ['name', 'num', 'date', 'num', 'num']
table = ttk.Treeview(frame_list, show='headings')
table['columns'] = heads

i = 0
for header in heads:
    table.heading(header, text=header, command=lambda _col=header: treeview_sort_column(table, _col, False),
                  anchor=anch[i])
    w_ = len(header) * 2 + 80
    table.column(header, width=w_, anchor=anch[i])
    i += 1
scroll_pane = ttk.Scrollbar(frame_list, command=table.yview)
table.configure(yscrollcommand=scroll_pane.set)
scroll_pane.pack(side=RIGHT, fill=Y)
table.pack(expand=YES, fill=BOTH)

frame_button_table = Frame(root, bg='black', bd=3)
frame_button_table.place(relx=0.05, rely=0.75, relwidth=0.2, relheight=0.05)
frame_button_table_num = Frame(root, bg='black', bd=3)
frame_button_table_num.place(relx=0.05, rely=0.7, relwidth=0.2, relheight=0.05)
frame_button_table_num_text = Frame(root, bg='black', bd=3)
frame_button_table_num_text.place(relx=0.05, rely=0.65, relwidth=0.2, relheight=0.05)
info = Label(frame_button_table_num_text, text='Введите количество записей', bg='black', font=40)
info.pack()
numField = Entry(frame_button_table_num, bg='gray', font=30)
numField.pack()
btn = Button(frame_button_table, text='Заполнить таблицу', command=get_table)
btn.pack()

frame_main = Frame(root, bg='black', bd=3)
frame_main.place(relx=0.05, rely=0.85, relwidth=0.2, relheight=0.05)
btn2 = Button(frame_main, text='Заполнить документ', command=func1)
btn2.pack()

frame_open = Frame(root, bg='black', bd=3)
frame_open.place(relx=0.05, rely=0.8, relwidth=0.2, relheight=0.05)
btn3 = Button(frame_open, text='Выбрать шаблон', command=func2)
btn3.pack()

frame_individual = Frame(root, bg='black', bd=3)
frame_individual.place(relx=0.025, rely=0.9, relwidth=0.25, relheight=0.05)
btn4 = Button(frame_individual, text='Заменить в тексте все точки на запятые', command=func3)
btn4.pack()

root.mainloop()
